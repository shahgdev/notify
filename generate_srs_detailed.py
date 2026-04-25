from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_detailed_srs():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- Title Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Software Requirements Specification\nFor\n<NOTIFY: AI-Based Lecture Recorder & Analyzer>')
    run.bold = True
    font = run.font
    font.size = Pt(24)
    font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('\n' * 3)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Version 1.0')
    run.font.size = Pt(14)
    
    doc.add_paragraph('\n' * 2)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Prepared by <M. Ibrahim Khan>\nReg No: B24F2174AI075\n<Abdullah Siddique>\nReg No: B24F2089AI100')
    run.font.size = Pt(12)
    
    doc.add_paragraph('\n' * 1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted To:\nMr. Musadaq Mansoor')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Date: March 08, 2026')
    
    doc.add_page_break()

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', 1)
    
    doc.add_heading('1.1 Purpose', 2)
    doc.add_paragraph('The purpose of this document is to provide a detailed description of the software requirements for the NOTIFY: AI-Based Lecture Recorder & Analyzer system. This SRS defines the functional and non-functional requirements to ensure that the system operates effectively for both students and educators. It ensures that all participants have a clear understanding of system objectives, capabilities, and constraints.')

    doc.add_heading('1.2 Document Conventions', 2)
    doc.add_paragraph('Standard IEEE formatting is followed. Requirements are identified as REQ-1, REQ-2, etc. Priority levels are assigned to each requirement.')

    doc.add_heading('1.3 Intended Audience and Reading Suggestions', 2)
    doc.add_paragraph('Intended for developers (technical reference), testers (designing test cases), and project managers (monitoring progress).')

    doc.add_heading('1.4 Product Scope', 2)
    doc.add_paragraph('NOTIFY helps students capture, transcribe, and analyze academic lectures. Users can record or upload audio files (MP3, WAV, etc.), which the system processes using AI to extract transcripts, generate summaries, interactive flashcards, and quizzes. It aims to reduce technical barriers in learning and improve information retention.')

    doc.add_page_break()

    # --- 2. Overall Description ---
    doc.add_heading('2. Overall Description', 1)
    
    doc.add_heading('2.1 Product Perspective', 2)
    doc.add_paragraph('NOTIFY is a standalone web-based platform designed to assist students in analyzing lectures and identifying key concepts. The system functions as an intelligent assistant where users upload audio and receive automated pedagogical feedback.')

    doc.add_heading('2.2 Product Functions', 2)
    doc.add_paragraph('- Audio Recording/Uploading\n- Speech-to-Text Transcription\n- Automated Summary Generation\n- Flashcard/Quiz Creation\n- Subject-based Organization\n- Multi-language Reporting')

    doc.add_heading('2.3 User Classes and Characteristics', 2)
    doc.add_paragraph('- Students: Primary users focused on learning from lectures.\n- Educators: Users who may upload their own lectures to provide structured materials.\n- Administrators: System managers responsible for maintenance and database updates.')

    doc.add_heading('2.4 Operating Environment', 2)
    doc.add_paragraph('Web-based accessible via modern browsers (Chrome, Edge, Firefox). Backend runs on FastAPI (Python) with SQLite/PostgreSQL. Frontend uses Vite, React, and Tailwind CSS.')

    doc.add_heading('2.5 Design and Implementation Constraints', 2)
    doc.add_paragraph('- Supports MP3, WAV, M4A, and high-quality recording formats.\n- Database integration (SQLAlchemy) for persistence.\n- Multi-browser compatibility.\n- Modern UI/UX standards (shadcn-ui).')

    doc.add_page_break()

    # --- 3. External Interface Requirements ---
    doc.add_heading('3. External Interface Requirements', 1)
    doc.add_heading('3.1 User Interfaces', 2)
    doc.add_paragraph('Web-based dashboard built with React. Features include drag-and-drop uploading, interactive transcript view, and summary/flashcard card views. Mobile responsive design is mandatory.')

    doc.add_heading('3.3 Software Interfaces', 2)
    doc.add_paragraph('- OS: Windows 10/11, Linux, macOS.\n- Database: PostgreSQL/SQLite.\n- AI Engine: OpenAI Whisper (Transcription) and Gemini/OpenRouter (Synthesis).')

    doc.add_page_break()

    # --- 4. System Features ---
    doc.add_heading('4. System Features', 1)
    
    doc.add_heading('4.1 Lecture Upload & Transcription', 2)
    doc.add_paragraph('REQ-1: Validates audio file duration and size (max 100MB).\nREQ-2: Extracts high-accuracy transcripts from uploads.\nREQ-3: Supports real-time recording via browser microphone API.')

    doc.add_heading('4.2 AI Pedagogical Analysis', 2)
    doc.add_paragraph('REQ-4: Generates bulleted summaries in selected languages.\nREQ-5: Extracts terms/definitions for Flashcards.\nREQ-6: Creates Multiple Choice Questions (MCQs) for self-testing.\nREQ-7: Generates AI-Topic Tags (e.g., [Quantum Mechanics]).')

    doc.add_heading('4.3 Subject Management', 2)
    doc.add_paragraph('REQ-8: Users can create/rename custom subjects.\nREQ-9: Lectures can be moved between subjects easily.')

    doc.add_page_break()

    # --- 5. Nonfunctional Requirements ---
    doc.add_heading('5. Other Nonfunctional Requirements', 1)
    
    doc.add_heading('5.1 Performance Requirements', 2)
    doc.add_paragraph('Transcription speed: < 5 minutes for 1-hour audio. UI response: < 200ms for navigation.')

    doc.add_heading('5.3 Security Requirements', 2)
    doc.add_paragraph('Encrypted database storages. API key rotation for cloud services. User data isolation.')

    doc.add_page_break()

    # --- Appendices ---
    doc.add_heading('Appendix A: Glossary', 1)
    doc.add_paragraph('AI Analyzer: The module using NLP to process transcripts.\nFlashcard: A study card with a term and definition.\nTranscription: Converting audio speech to text.')

    # Save
    doc.save('NOTIFY_SRS_Detailed.docx')
    print("Detailed SRS Document 'NOTIFY_SRS_Detailed.docx' created successfully.")

if __name__ == "__main__":
    create_detailed_srs()
