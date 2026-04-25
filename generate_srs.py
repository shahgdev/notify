from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_srs():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('Software Requirements Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('For', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('<NOTIFY: AI-Based Lecture Recorder & Analyzer>', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    
    version = doc.add_paragraph('Version 1.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    prepared_by = doc.add_paragraph('Prepared by <Drafted by Antigravity AI>')
    prepared_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph('Date: March 08, 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('1. Introduction\n2. Overall Description\n3. External Interface Requirements\n4. System Features\n5. Other Nonfunctional Requirements\n6. Other Requirements\nAppendix A: Glossary')
    doc.add_page_break()

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', 1)
    
    doc.add_heading('1.1 Purpose', 2)
    doc.add_paragraph('The purpose of this document is to provide a detailed description of the software requirements for the NOTIFY: AI-Based Lecture Recorder & Analyzer system. This SRS defines the functional and non-functional requirements to ensure effective system operation for students and educators.')

    doc.add_heading('1.2 Document Conventions', 2)
    doc.add_paragraph('Standard IEEE formatting is followed. Requirements are identified as REQ-1, REQ-2, etc.')

    doc.add_heading('1.4 Product Scope', 2)
    doc.add_paragraph('NOTIFY helps students capture, transcribe, and analyze academic lectures. Users can record or upload audio files, which the system then processes to extract transcripts, generate concise summaries, create flashcards, and build quizzes. It also provides subject-based organization and AI-generated topic tagging.')

    # --- 2. Overall Description ---
    doc.add_heading('2. Overall Description', 1)
    
    doc.add_heading('2.1 Product Perspective', 2)
    doc.add_paragraph('NOTIFY is a standalone web-based platform designed to enhance the learning experience through automated content analysis.')

    doc.add_heading('2.2 Product Functions', 2)
    doc.add_paragraph('- Audio Recording and Uploading\n- Speech-to-Text Transcription\n- AI Summary Generation\n- Flashcard and Quiz Creation\n- Subject Management\n- Multi-language Support (English, Urdu, etc.)')

    doc.add_heading('2.4 Operating Environment', 2)
    doc.add_paragraph('Web-based accessible via Chrome, Firefox, and Edge. Backend runs on FastAPI with a SQLite/PostgreSQL database.')

    # --- 3. External Interface Requirements ---
    doc.add_heading('3.External Interface Requirements', 1)
    doc.add_heading('3.1 User Interfaces', 2)
    doc.add_paragraph('A modern, responsive dashboard built with React and Tailwind CSS. Key screens: Lecture Dashboard, Recorder, Subject Manager, and Analytics View.')

    doc.add_heading('3.3 Software Interfaces', 2)
    doc.add_paragraph('- OS: Windows, Linux, macOS\n- Database: SQLite\n- AI: OpenRouter / OpenAI API\n- Transcription: Faster-Whisper / OpenAI Whisper')

    # --- 4. System Features ---
    doc.add_heading('4. System Features', 1)
    
    doc.add_heading('4.1 Lecture Upload & Processing', 2)
    doc.add_paragraph('REQ-1: Accept MP3, WAV, M4A formats.\nREQ-2: Extract audio from video uploads (optional).\nREQ-3: Store files securely in local/cloud storage.')

    doc.add_heading('4.2 AI Analysis & Content Generation', 2)
    doc.add_paragraph('REQ-4: Generate accurate transcriptions.\nREQ-5: Produce structured summaries in multiple languages.\nREQ-6: Generate interactive flashcards and quizzes.\nREQ-7: Assign AI-generated topic tags for organization.')

    # --- 5. Other Nonfunctional Requirements ---
    doc.add_heading('5. Other Nonfunctional Requirements', 1)
    
    doc.add_heading('5.1 Performance Requirements', 2)
    doc.add_paragraph('Transcription and analysis should complete within 2-5 minutes for a 60-minute lecture.')

    doc.add_heading('5.3 Security Requirements', 2)
    doc.add_paragraph('Data at rest encrypted using AES-256. Secure API communication via HTTPS.')

    # --- Save ---
    doc.save('NOTIFY_SRS.docx')
    print("SRS Document 'NOTIFY_SRS.docx' created successfully.")

if __name__ == "__main__":
    create_srs()
